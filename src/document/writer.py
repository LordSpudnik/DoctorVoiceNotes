"""
writer.py
=========

Writes the doctor's dictated notes into a Microsoft Word (.docx) document
on disk, per PRD FR-05 (Automatic Save), FR-06 (Word Document Generation),
FR-07 (Stop Recording final save), and Section 11 (Document Format).

This module is the consumer of src.commands.voice_commands's
VoiceCommandProcessor. It knows nothing about audio, VAD, or Whisper - it
only calls processor.get_pending_text() / mark_saved() / has_pending_text()
/ consume_save_request(), and turns whatever text comes back into
python-docx paragraphs. See the Phase 4 handoff brief, Section 5, for the
open interface question this module exists to resolve.

------------------------------------------------------------------------
THE INTERFACE QUESTION THIS MODULE RESOLVES (read this first)
------------------------------------------------------------------------
VoiceCommandProcessor.get_pending_text() renders every segment recorded
since the last mark_saved() call, INCLUDING a trailing in-progress
sentence that has no closing "full stop" yet (this is deliberate -
FR-05/FR-07 crash safety: an unfinished sentence must never simply
vanish if the app dies mid-dictation). Concretely, across two consecutive
autosave cycles you can see:

    cycle 1: get_pending_text() -> "Patient complains of chest pain"
    cycle 2: get_pending_text() -> "for two days."

Both of those are really the SAME sentence - the doctor kept talking
between cycle 1 and cycle 2, and "full stop" only landed in cycle 2. If
this writer naively appended each cycle's text as a brand new docx
paragraph, that one sentence would be split across two separate lines in
the Word document, which violates PRD Section 11's "one sentence per
line" format.

THE FIX: this writer never needs VoiceCommandProcessor to change (Phase 4
is done and user-confirmed; this module works entirely from the plain
text strings the existing interface already returns). It works because
of an invariant in render_segments() (voice_commands.py) that is a fact
about the text, not something this module has to guess at:

    Within any single get_pending_text() batch, split on "\n":
      - Every line EXCEPT POSSIBLY THE LAST was closed by an explicit
        boundary command (full stop / new paragraph / new patient) and
        will NEVER be extended by a future batch.
      - The LAST line is still "open" (may be extended by the next
        batch's first line) if and only if it is non-empty AND does not
        end with "." - because a "." can only appear via the "full stop"
        command (Whisper's own punctuation is stripped before this
        writer ever sees the text), and an empty last line means the
        batch ended exactly on a paragraph-break boundary with nothing
        left dangling.

So this writer tracks a single piece of state, self._pending_paragraph:
a reference to the python-docx Paragraph object that is currently "open"
(or None if the last thing written was closed). Each new batch's FIRST
line, if there is an open paragraph to extend, is appended to it via
add_run() instead of starting a new paragraph. Every other line starts a
fresh paragraph. See _write_pending_text() for the exact algorithm.

ASSUMPTIONS FLAGGED IN THIS PHASE (not confirmed by user - see chat reply)
------------------------------------------------------------------------
1. PRD Section 11 shows a Date/Time header ("29 July 2026" / "10:45 AM")
   at the top of the document, directly followed by note content - no
   literal word "Patient" (that literal word is specific to the "new
   patient" VOICE COMMAND from Phase 4, a different mechanism). This
   writer assumes that header is written once per RECORDING SESSION
   (i.e. once per Start button press - start_session()), not once ever
   per document file. The PRD does not say which, and it matters: if the
   doctor Starts/Stops multiple times against the same PatientNotes.docx
   in one day, this assumption produces a repeated Date/Time block each
   time. Flagging this explicitly; easy to change to "only on first-ever
   document creation" if that is not what is wanted.
2. If start_session() opens an EXISTING .docx (FR-06: "append if the
   document already exists"), self._pending_paragraph always starts as
   None for that session - i.e. new dictation is never spliced onto
   whatever paragraph was last in the file, even if that file's last
   session ended mid-sentence (e.g. a crash, or force-quit, before a
   proper Stop). This is the conservative choice: silently re-opening
   and appending onto a stale paragraph from a possibly-earlier patient
   or a possibly-already-read note is a worse failure mode than leaving
   one incomplete-looking line in the document. Flagging as a known
   limitation, not solved further here.
3. No additional docx styling (fonts, bold headers, colours) is applied
   beyond plain paragraphs - the PRD's "neatly formatted" secondary
   objective is read as "one sentence per line, blank lines around
   headers", which is what Section 11's own example shows. Cosmetic
   styling is left as an easy follow-up if the doctor wants it.
4. Rollback-on-failed-save (see _rollback()) reaches into python-docx's
   private `_element` attribute to delete a paragraph/run, because the
   public python-docx API has no supported "delete paragraph" method.
   This is a widely-used community workaround, not something invented
   here, but it is technically an internal-API dependency - flagging the
   risk that a future python-docx upgrade could change this. requirements
   .txt should keep python-docx version-pinned (should already be pinned
   from Phase 1) to avoid surprises.
"""

import shutil
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from src.utils.config_manager import ConfigManager
from src.utils.logger import get_logger
from src.utils.paths import get_app_root

logger = get_logger(__name__)


# Retry policy for FR-05/Section 14 ("Cannot save document -> Retry
# automatically"). Kept short: this covers brief transient locks (e.g. a
# cloud-sync client or antivirus scanner briefly holding the file), not
# a long outage - a doctor mid-dictation should not be blocked for long.
_MAX_SAVE_ATTEMPTS = 3
_RETRY_DELAY_SECONDS = 1.0

# How often run_autosave_loop() wakes up to check for a "save note"
# command, independent of the doctor-configured autosave_interval_seconds.
# Short enough that "save note" feels immediate (FR-08), long enough to
# not busy-loop.
_LOOP_POLL_SECONDS = 0.25


class DocumentWriter:
    """
    Owns the on-disk .docx file for the current recording session and the
    in-memory python-docx Document object backing it.

    Typical usage (Phase 6/7 wiring):

        writer = DocumentWriter(config, processor)
        writer.start_session()                     # on Start (FR-02)
        # ... in a background thread, mirroring engine.py's
        # transcribe_stream() pattern (this class does not spawn its own
        # thread - see run_autosave_loop() docstring):
        writer.run_autosave_loop(stop_event, config.get("autosave_interval_seconds"))
        # ... on Stop (FR-07):
        writer.stop_session()
    """

    def __init__(
        self,
        config: ConfigManager,
        processor,
        clock: Callable[[], datetime] = datetime.now,
    ):
        """
        Args:
            config: the app's ConfigManager (already loaded), for
                save-folder/document-name settings (FR-09).
            processor: a VoiceCommandProcessor instance (Phase 4). Typed
                loosely (duck-typed) rather than imported directly as a
                hard type-hint dependency, to keep this module easy to
                unit test against a minimal fake if ever needed.
            clock: injectable clock, defaulting to real datetime.now -
                same dependency-injection pattern used by
                VoiceCommandProcessor and PhraseSegmenter's tests.
        """
        self._config = config
        self._processor = processor
        self._clock = clock

        self._doc: Optional[Document] = None
        self._doc_path: Optional[Path] = None
        self._pending_paragraph: Optional[Paragraph] = None

    # ------------------------------------------------------------------
    # SESSION LIFECYCLE
    # ------------------------------------------------------------------
    def start_session(self) -> None:
        """
        Call when the doctor presses Start (FR-02). Resolves the save
        path from current settings, opens the existing document or
        creates a new one (FR-06), writes this session's Date/Time header
        (Section 11), and saves immediately so the file exists on disk
        from the first moment of recording (not only once the first
        autosave tick fires).
        """
        self._doc_path = self._resolve_document_path()

        if self._doc_path.exists():
            try:
                self._doc = Document(str(self._doc_path))
                logger.info(f"Opened existing document for appending: {self._doc_path}")
            except Exception as e:
                # Mirrors ConfigManager's corrupted-settings.json handling
                # (Section 14: prevent crash whenever possible). Broad
                # `except Exception` is deliberate here - python-docx does
                # not expose one single documented exception type that
                # covers every way a .docx file can be unreadable
                # (missing zip entries, bad XML, wrong file entirely).
                logger.error(
                    f"Existing document at {self._doc_path} could not be "
                    f"opened ({e}). Backing it up and starting a fresh "
                    f"document so the session can still proceed."
                )
                self._backup_unreadable_document()
                self._doc = Document()
        else:
            self._doc = Document()
            logger.info(f"Creating new document: {self._doc_path}")

        # A fresh session never inherits an "open" paragraph from
        # whatever was last in the file - see module docstring,
        # assumption 2.
        self._pending_paragraph = None

        self._write_session_header()
        if not self._save_to_disk():
            logger.error(
                f"Could not write the initial document to {self._doc_path} "
                f"at session start - will keep retrying on the next "
                f"autosave cycle."
            )

    def autosave(self) -> bool:
        """
        One save cycle (FR-05). Pulls whatever is new from the voice
        command processor, writes it into the docx, and flushes to disk.

        Returns:
            True if there was new text AND it was successfully written
            and saved. False if there was nothing new to save (FR-05:
            "save only newly recognised text" - this is the normal,
            expected case on most autosave ticks, not an error), or if
            writing/saving failed after retries (in which case nothing
            is left half-written - see _rollback()).
        """
        if self._doc is None:
            raise RuntimeError("autosave() called before start_session().")

        # Deliberately checking has_pending_text() rather than truthiness
        # of get_pending_text() itself: a lone "new paragraph"/"new
        # patient" command with no dictated text around it yet renders to
        # "" (an empty string), but it is still real pending content (a
        # blank-line boundary) that must eventually be written and marked
        # saved - otherwise it would sit in limbo forever and, worse,
        # stop_session() would misreport a completely benign "nothing to
        # write yet" as a save FAILURE. _write_pending_text("") correctly
        # writes a single blank paragraph for this case.
        if not self._processor.has_pending_text():
            return False

        pending_text = self._processor.get_pending_text()
        pending_paragraph_before = self._pending_paragraph
        mutations = self._write_pending_text(pending_text)

        if self._save_to_disk():
            self._processor.mark_saved()
            logger.info("Autosave: wrote and saved new content.")
            return True

        # Disk save failed after retries - undo the in-memory paragraph
        # mutations so the NEXT autosave cycle (which will see the same
        # pending text again, since mark_saved() was never called) does
        # not duplicate what we just tried to add.
        self._rollback(mutations, pending_paragraph_before)
        logger.error(
            "Autosave: could not persist to disk after retries. Rolled "
            "back in-memory changes; this text remains pending and will "
            "be retried on the next autosave cycle."
        )
        return False

    def stop_session(self) -> None:
        """
        Call when the doctor presses Stop (FR-07): completes
        transcription is assumed to already have happened upstream
        (engine.py's transcribe_stream() drains remaining audio and
        pushes a final sentinel); this method's job is to make sure
        whatever text that produced gets written and saved before the
        app reports Idle.
        """
        if self._doc is None:
            logger.warning("stop_session() called before start_session(); nothing to do.")
            return

        if self._processor.has_pending_text():
            success = self.autosave()
            if not success:
                # This is worse than a mid-session autosave miss: FR-07
                # promises "Word document is saved" as part of Stop
                # completing. Phase 6/7's UI must surface this to the
                # doctor rather than silently reporting Idle - flagging
                # as a cross-phase dependency, not solved here.
                logger.error(
                    "FR-07 final save on Stop FAILED after retries. "
                    "Dictated text exists only in memory. The UI layer "
                    "must warn the doctor rather than silently return to "
                    "Idle."
                )
        else:
            # Nothing new, but make sure the file on disk reflects the
            # in-memory document (should already be in sync from the
            # last successful autosave tick - this is a defensive no-op
            # in the common case).
            self._save_to_disk()

        logger.info("Recording session stopped; final save attempted.")

    # ------------------------------------------------------------------
    # AUTOSAVE LOOP
    # ------------------------------------------------------------------
    def run_autosave_loop(self, stop_event: threading.Event, interval_seconds: float) -> None:
        """
        BLOCKING loop - run this in a dedicated background thread. This
        class does not spawn its own thread; the caller (Phase 6 UI /
        Phase 7 main.py) is responsible for that - the same convention
        TranscriptionEngine.transcribe_stream() uses (see engine.py's
        module docstring), which keeps this class simple to unit test
        (tests call autosave() directly, with no real threading or
        timing involved).

        Wakes up every _LOOP_POLL_SECONDS to check whether "save note"
        was heard (processor.consume_save_request()) - so an immediate
        save request doesn't sit waiting out the rest of a longer
        interval - and otherwise fires a normal autosave every
        `interval_seconds`. Returns (and the caller's thread exits)
        once `stop_event` is set.
        """
        elapsed = 0.0
        while not stop_event.is_set():
            if stop_event.wait(timeout=_LOOP_POLL_SECONDS):
                break
            elapsed += _LOOP_POLL_SECONDS
            if self._processor.consume_save_request() or elapsed >= interval_seconds:
                elapsed = 0.0
                self.autosave()
        logger.info("Autosave loop stopped.")

    # ------------------------------------------------------------------
    # PATH RESOLUTION (FR-09 settings)
    # ------------------------------------------------------------------
    def _resolve_document_path(self) -> Path:
        """
        Reads default_save_folder / default_document_name from config
        FRESH each time start_session() runs (not cached at construction)
        so a settings-dialog change (FR-09) before the next Start takes
        effect. A relative folder is resolved under the app root, per
        paths.py's "persistent data always lives beside the exe/project
        root, never the process's cwd" rule; an absolute folder (e.g. a
        doctor-chosen custom location) is used as-is.
        """
        folder_setting = Path(self._config.get("default_save_folder", "notes"))
        folder = folder_setting if folder_setting.is_absolute() else get_app_root() / folder_setting
        folder.mkdir(parents=True, exist_ok=True)

        doc_name = self._config.get("default_document_name", "PatientNotes.docx")
        return folder / doc_name

    # ------------------------------------------------------------------
    # SESSION HEADER (Section 11)
    # ------------------------------------------------------------------
    def _write_session_header(self) -> None:
        """
        Writes the Date/Time header shown at the top of PRD Section 11's
        worked example. See module docstring, assumption 1, for why this
        runs once per session rather than once per document.
        """
        now = self._clock()
        date_str = now.strftime("%d %B %Y")
        time_str = now.strftime("%I:%M %p")
        self._doc.add_paragraph(date_str)
        self._doc.add_paragraph(time_str)
        self._doc.add_paragraph("")  # blank line before notes (Sec. 11 example)
        logger.info(f"Session header written: {date_str} {time_str}")

    # ------------------------------------------------------------------
    # CORE WRITE ALGORITHM (see module docstring for why this is correct)
    # ------------------------------------------------------------------
    def _write_pending_text(self, text: str) -> list[tuple]:
        """
        Writes `text` (as returned by processor.get_pending_text()) into
        self._doc, extending the still-open paragraph from a previous
        cycle where applicable instead of duplicating it.

        Returns the list of mutations made, so autosave() can precisely
        undo them via _rollback() if the subsequent disk save fails.
        Each mutation is one of:
            ("new_paragraph", paragraph)
            ("extend_run", paragraph, run)
        """
        mutations: list[tuple] = []
        lines = text.split("\n")
        last_index = len(lines) - 1

        for i, line in enumerate(lines):
            if line == "":
                # A blank line is always a closed boundary (paragraph
                # break / patient header spacing) - never something to
                # extend into, and never itself left "open".
                paragraph = self._doc.add_paragraph("")
                mutations.append(("new_paragraph", paragraph))
                self._pending_paragraph = None
                continue

            if i == 0 and self._pending_paragraph is not None:
                # Continuation of the paragraph left open by a previous
                # autosave cycle - the core fix this module exists for.
                run_text = f" {line}" if self._pending_paragraph.text else line
                run = self._pending_paragraph.add_run(run_text)
                mutations.append(("extend_run", self._pending_paragraph, run))
                paragraph = self._pending_paragraph
            else:
                paragraph = self._doc.add_paragraph(line)
                mutations.append(("new_paragraph", paragraph))

            # Only the LAST line of a batch can still be "open" - every
            # earlier line was necessarily closed by a boundary command
            # for it to no longer be the last line (see module docstring).
            if i == last_index and not line.endswith("."):
                self._pending_paragraph = paragraph
            else:
                self._pending_paragraph = None

        return mutations

    def _rollback(self, mutations: list[tuple], pending_paragraph_before: Optional[Paragraph]) -> None:
        """
        Undoes exactly the mutations _write_pending_text() just made, so
        a failed disk save leaves self._doc exactly as it was before this
        autosave attempt - see module docstring, assumption 4, for the
        internal-API caveat this relies on.
        """
        for mutation in reversed(mutations):
            if mutation[0] == "new_paragraph":
                paragraph = mutation[1]
                paragraph._element.getparent().remove(paragraph._element)
            elif mutation[0] == "extend_run":
                _, _paragraph, run = mutation
                run._element.getparent().remove(run._element)
        self._pending_paragraph = pending_paragraph_before

    # ------------------------------------------------------------------
    # DISK I/O
    # ------------------------------------------------------------------
    def _save_to_disk(self) -> bool:
        """
        Saves self._doc to self._doc_path, retrying a few times on
        failure (Section 14: "Cannot save document -> Retry
        automatically") before giving up. Never raises - callers get a
        bool and decide what that means for the processor's saved state.
        """
        for attempt in range(1, _MAX_SAVE_ATTEMPTS + 1):
            try:
                self._doc.save(str(self._doc_path))
                return True
            except OSError as e:
                logger.error(
                    f"Save attempt {attempt}/{_MAX_SAVE_ATTEMPTS} failed "
                    f"for {self._doc_path}: {e}"
                )
                if attempt < _MAX_SAVE_ATTEMPTS:
                    time.sleep(_RETRY_DELAY_SECONDS)
        return False

    def _backup_unreadable_document(self) -> None:
        """Renames an unreadable existing .docx to
        <name>.docx.broken-<timestamp> instead of silently overwriting
        it, mirroring ConfigManager._backup_corrupted_file()."""
        timestamp = self._clock().strftime("%Y%m%d-%H%M%S")
        backup_path = self._doc_path.with_suffix(f".docx.broken-{timestamp}")
        try:
            shutil.copy2(self._doc_path, backup_path)
            logger.warning(f"Unreadable document backed up to {backup_path}")
        except OSError as e:
            logger.error(f"Could not back up the unreadable document: {e}")